from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.contrib import messages
from django.db.models import Count, Q
from django.utils import timezone
from django.utils.html import escape
from .models import Report, ReportTemplate
from worklist.models import Study
from accounts.models import User
import io
from django.core.files.base import ContentFile
from django.utils.text import slugify
from django.urls import reverse
from django.conf import settings
import io
import base64

# Optional PDF/Docx libs
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None
try:
    from docx import Document
except Exception:
    Document = None

# QR code support
try:
    import qrcode
except Exception:
    qrcode = None


def _qr_png_bytes(url: str, box_size: int = 6, border: int = 2) -> bytes:
    """Return PNG bytes for a QR code for the given URL."""
    if not qrcode:
        return b''
    qr = qrcode.QRCode(
        version=None,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=box_size,
        border=border,
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color='black', back_color='white').convert('RGB')
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    return buf.getvalue()


def _data_url_from_png(png_bytes: bytes) -> str:
    if not png_bytes:
        return ''
    return 'data:image/png;base64,' + base64.b64encode(png_bytes).decode('ascii')


@login_required
def report_list(request):
    # Restrict to admin and radiologist
    if not getattr(request.user, 'can_edit_reports', None) or not request.user.can_edit_reports():
        return HttpResponse(status=403)
    """List all reports"""
    # Get filter parameters
    search_query = request.GET.get('search', '')
    status_filter = request.GET.get('status', '')
    modality_filter = request.GET.get('modality', '')
    date_from = request.GET.get('date_from', '')
    date_to = request.GET.get('date_to', '')
    sort = request.GET.get('sort', '')
    order = request.GET.get('order', 'desc')

    # Base queryset
    reports = Report.objects.select_related('study', 'study__patient', 'study__modality', 'radiologist').all()

    # Apply filters
    if search_query:
        reports = reports.filter(
            Q(study__patient__first_name__icontains=search_query) |
            Q(study__patient__last_name__icontains=search_query) |
            Q(study__accession_number__icontains=search_query) |
            Q(radiologist__first_name__icontains=search_query) |
            Q(radiologist__last_name__icontains=search_query)
        )

    if status_filter:
        reports = reports.filter(status=status_filter)

    if modality_filter:
        reports = reports.filter(study__modality__code=modality_filter)

    if date_from:
        try:
            reports = reports.filter(report_date__date__gte=date_from)
        except Exception:
            pass

    if date_to:
        try:
            reports = reports.filter(report_date__date__lte=date_to)
        except Exception:
            pass

    # Sorting
    sort_fields = {
        'patient': 'study__patient__last_name',
        'study_date': 'study__study_date',
        'radiologist': 'radiologist__last_name',
        'report_date': 'report_date',
    }
    sort_field = sort_fields.get(sort, 'report_date')
    if order == 'asc':
        reports = reports.order_by(sort_field)
    else:
        reports = reports.order_by(f'-{sort_field}')

    # Calculate statistics
    total_reports = Report.objects.count()
    draft_reports = Report.objects.filter(status='draft').count()
    pending_reports = Report.objects.filter(status='preliminary').count()
    final_reports = Report.objects.filter(status='final').count()
    amended_reports = Report.objects.filter(status='amended').count()
    cancelled_reports = Report.objects.filter(status='cancelled').count()

    context = {
        'reports': reports,
        'total_reports': total_reports,
        'draft_reports': draft_reports,
        'pending_reports': pending_reports,
        'final_reports': final_reports,
        'amended_reports': amended_reports,
        'cancelled_reports': cancelled_reports,
        'search_query': search_query,
        'status_filter': status_filter,
        'modality_filter': modality_filter,
        'date_from': date_from,
        'date_to': date_to,
        'sort': sort,
        'order': order,
    }

    return render(request, 'reports/report_list.html', context)


@login_required
def write_report(request, study_id):
    # Restrict to admin and radiologist
    if not getattr(request.user, 'can_edit_reports', None) or not request.user.can_edit_reports():
        return HttpResponse(status=403)
    """Write report for study"""
    study = get_object_or_404(Study, id=study_id)
    user = request.user

    # Radiologists/admins scoped to their assigned facility
    if not user.is_admin() and hasattr(user, 'facility') and user.facility:
        if study.facility and study.facility != user.facility:
            from django.contrib import messages as _msgs
            _msgs.error(request, 'You can only write reports for studies at your assigned facility.')
            return redirect('worklist:dashboard')

    # When opening editor, mark study as in progress for editors
    if study.status in ['scheduled', 'suspended']:
        study.status = 'in_progress'
        study.save(update_fields=['status'])
    
    # Try to get existing report or create new one
    try:
        report = Report.objects.get(study=study)
        is_new_report = False
    except Report.DoesNotExist:
        report = None
        is_new_report = True
    
    if request.method == 'POST':
        # Get form data
        clinical_history = request.POST.get('clinical_history', '')
        technique = request.POST.get('technique', '')
        comparison = request.POST.get('comparison', '')
        findings = request.POST.get('findings', '')
        impression = request.POST.get('impression', '')
        recommendations = request.POST.get('recommendations', '')
        status = request.POST.get('status', 'draft')
        action = request.POST.get('action', 'save')
        
        if is_new_report:
            # Create new report
            report = Report.objects.create(
                study=study,
                radiologist=request.user,
                clinical_history=clinical_history,
                technique=technique,
                comparison=comparison,
                findings=findings,
                impression=impression,
                recommendations=recommendations,
                status=status
            )
            messages.success(request, 'Report created successfully!')
        else:
            # Update existing report
            report.clinical_history = clinical_history
            report.technique = technique
            report.comparison = comparison
            report.findings = findings
            report.impression = impression
            report.recommendations = recommendations
            report.status = status
            report.last_modified = timezone.now()
            
            # If finalizing the report
            if status == 'final' or (action == 'submit' and status == 'final'):
                report.signed_date = timezone.now()
            
            report.save()
            messages.success(request, 'Report updated successfully!')
        
        # Update study status when report finalized
        if status == 'final' or (action == 'submit' and status == 'final'):
            study.status = 'completed'
            study.save(update_fields=['status'])
        
        # Redirect based on action
        if action == 'submit':
            messages.success(request, 'Report submitted successfully!')
            return redirect('reports:report_list')
        else:
            # Stay on the same page for continued editing
            return redirect('reports:write_report', study_id=study_id)
    
    context = {
        'study': study,
        'report': report,
        'is_new_report': is_new_report,
    }
    
    return render(request, 'reports/write_report.html', context)


@login_required
def print_report_stub(request, study_id):
    """Printable HTML that mirrors facility letterhead, includes author signature and QR/link footer."""
    study = get_object_or_404(Study, id=study_id)

    # Facility users can only print their own facility's reports
    user = request.user
    if not user.can_edit_reports():
        if not hasattr(user, 'facility') or not user.facility or study.facility != user.facility:
            from django.http import HttpResponseForbidden
            return HttpResponseForbidden('Access denied')

    report = Report.objects.filter(study=study).first()

    # Build absolute URLs
    viewer_url = request.build_absolute_uri(reverse('dicom_viewer:web_viewer')) + f"?study_id={study.id}"
    report_url = request.build_absolute_uri(reverse('reports:print_report', args=[study.id]))

    # Generate QR codes
    qr_viewer_b64 = _data_url_from_png(_qr_png_bytes(viewer_url))
    qr_report_b64 = _data_url_from_png(_qr_png_bytes(report_url))

    # Embed letterhead directly to avoid relying on public /media/ URLs.
    # This keeps the printable report working even when MEDIA is not served publicly.
    letterhead_url = ''
    try:
        lh = getattr(study.facility, 'letterhead', None)
        if lh and getattr(lh, 'name', ''):
            import mimetypes
            ctype, _ = mimetypes.guess_type(lh.name)
            ctype = ctype or 'image/png'
            with lh.open('rb') as f:
                b = f.read()
            letterhead_url = f"data:{ctype};base64," + base64.b64encode(b).decode('ascii')
    except Exception:
        letterhead_url = ''
    return render(request, 'reports/print_report.html', {
        'study': study,
        'report': report,
        'patient': study.patient,
        'facility': study.facility,
        'qr_viewer': qr_viewer_b64,
        'qr_report': qr_report_b64,
        'letterhead_b64': letterhead_url,
    })


@login_required
def export_report_pdf(request, study_id):
    # Restrict to admin and radiologist
    if not getattr(request.user, 'can_edit_reports', None) or not request.user.can_edit_reports():
        return HttpResponse(status=403)
    study = get_object_or_404(Study, id=study_id)
    report = Report.objects.filter(study=study).first()

    # Prepare absolute URLs
    viewer_url = request.build_absolute_uri(reverse('dicom_viewer:web_viewer')) + f"?study_id={study.id}"
    report_url = request.build_absolute_uri(reverse('reports:print_report', args=[study.id]))

    filename = f"report_{slugify(study.accession_number)}.pdf"
    if fitz is None:
        return JsonResponse({'error': 'PDF export not available (PyMuPDF missing).'}, status=500)
    try:
        doc = fitz.open()
        page = doc.new_page()
        margin = 36
        y = margin

        # Insert facility letterhead image if available
        try:
            if getattr(study.facility, 'letterhead', None) and study.facility.letterhead.name:
                try:
                    letterhead_data = study.facility.letterhead.open('rb').read()
                except Exception:
                    letterhead_data = None
                if letterhead_data:
                    rect = fitz.Rect(margin, y, page.rect.width - margin, y + 90)
                    page.insert_image(rect, stream=letterhead_data, keep_proportion=True)
                    y = rect.y1 + 12
                else:
                    raise Exception("letterhead unreadable")
            else:
                page.insert_text((margin, y), study.facility.name, fontsize=14, fontname="helv", fill=(0, 0, 0))
                y += 22
                page.insert_text((margin, y), study.facility.address or '', fontsize=10, fontname="helv", fill=(0, 0, 0))
                y += 18
        except Exception:
            # Fallback to text header if image fails
            page.insert_text((margin, y), study.facility.name, fontsize=14, fontname="helv", fill=(0, 0, 0))
            y += 22

        # Horizontal rule
        page.draw_line((margin, y), (page.rect.width - margin, y), color=(0, 0, 0), width=0.6)
        y += 10

        # Patient and study header
        header_lines = [
            f"Radiology Report",
            f"Patient: {study.patient.full_name} ({study.patient.patient_id})",
            f"Accession: {study.accession_number}    Modality: {study.modality.code}    Date: {study.study_date}",
            f"Priority: {study.priority.upper()}",
        ]
        for line in header_lines:
            page.insert_text((margin, y), line, fontsize=12 if line == 'Radiology Report' else 10, fontname="helv", fill=(0, 0, 0))
            y += 16 if line != 'Radiology Report' else 20

        y += 4
        # Sections
        def add_section(title: str, content: str):
            nonlocal y, page
            if y > page.rect.height - 180:
                page = doc.new_page(); y = margin
            page.insert_text((margin, y), title, fontsize=11, fontname="helv", fill=(0, 0, 0))
            y += 14
            # simple wrap
            max_width = page.rect.width - margin * 2
            words = (content or '-').split()
            line = ''
            while words:
                nxt = words.pop(0)
                test = (line + ' ' + nxt).strip()
                # crude width estimation (6px per char at 10pt)
                if len(test) * 6 > max_width:
                    page.insert_text((margin, y), line, fontsize=10, fontname="helv", fill=(0, 0, 0))
                    y += 13
                    line = nxt
                else:
                    line = test
            if line:
                page.insert_text((margin, y), line, fontsize=10, fontname="helv", fill=(0, 0, 0))
                y += 16
            y += 2

        add_section('Clinical History', (report.clinical_history if report else (study.clinical_info or '')) or '-')
        add_section('Technique', (report.technique if report else '') or '-')
        add_section('Comparison', (report.comparison if report else '') or '-')
        add_section('Findings', (report.findings if report else '') or '-')
        add_section('Impression', (report.impression if report else '') or '-')
        add_section('Recommendations', (report.recommendations if report else '') or '-')

        # Signature block
        author_name = ''
        author_license = ''
        if report and getattr(report, 'radiologist', None):
            try:
                author_name = report.radiologist.get_full_name() or report.radiologist.username
                author_license = getattr(report.radiologist, 'license_number', '') or ''
            except Exception:
                pass
        sig_line = f"Signed by: {author_name}{(' - ' + author_license) if author_license else ''}"
        if report and report.signed_date:
            sig_line += f" on {report.signed_date.strftime('%Y-%m-%d %H:%M')}"
        if y > page.rect.height - 150:
            page = doc.new_page(); y = margin
        page.insert_text((margin, y), sig_line, fontsize=10, fontname="helv", fill=(0, 0, 0))
        y += 14
        # Optional signature image
        if report and (report.digital_signature or '').strip():
            ds = report.digital_signature.strip()
            try:
                if ds.startswith('data:image'):
                    b64 = ds.split(',', 1)[1]
                else:
                    b64 = ds
                img_bytes = base64.b64decode(b64)
                rect = fitz.Rect(margin, y, margin + 180, y + 60)
                page.insert_image(rect, stream=img_bytes, keep_proportion=True)
                y = rect.y1 + 8
            except Exception:
                pass

        # Footer with QR codes
        if y < page.rect.height - 130:
            footer_y = page.rect.height - 130
        else:
            footer_y = y + 10
        # line
        page.draw_line((margin, footer_y), (page.rect.width - margin, footer_y), color=(0, 0, 0), width=0.6)
        footer_y += 8

        try:
            qr1 = _qr_png_bytes(viewer_url)
            qr2 = _qr_png_bytes(report_url)
            left_rect = fitz.Rect(margin, footer_y, margin + 110, footer_y + 110)
            right_rect = fitz.Rect(page.rect.width - margin - 110, footer_y, page.rect.width - margin, footer_y + 110)
            if qr1:
                page.insert_image(left_rect, stream=qr1, keep_proportion=True)
                page.insert_text((left_rect.x0, left_rect.y1 + 2), 'Scan to view images', fontsize=8, fontname='helv', fill=(0,0,0))
            if qr2:
                page.insert_image(right_rect, stream=qr2, keep_proportion=True)
                page.insert_text((right_rect.x0, right_rect.y1 + 2), 'Scan to view report', fontsize=8, fontname='helv', fill=(0,0,0))
        except Exception:
            pass

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        resp = HttpResponse(buf.getvalue(), content_type='application/pdf')
        resp['Content-Disposition'] = f'attachment; filename="{filename}"'
        return resp
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def export_report_docx(request, study_id):
    # Restrict to admin and radiologist
    if not getattr(request.user, 'can_edit_reports', None) or not request.user.can_edit_reports():
        return HttpResponse(status=403)
    if Document is None:
        return JsonResponse({'error': 'DOCX export not available (python-docx missing).'}, status=500)
    study = get_object_or_404(Study, id=study_id)
    report = Report.objects.filter(study=study).first()
    doc = Document()
    # Letterhead as image if available
    try:
        if getattr(study.facility, 'letterhead', None) and study.facility.letterhead.name:
            lh_data = study.facility.letterhead.open('rb').read()
            doc.add_picture(io.BytesIO(lh_data), width=None)
        else:
            raise Exception("no letterhead")
    except Exception:
        doc.add_heading(study.facility.name, 0)
    doc.add_paragraph(f"Patient: {study.patient.full_name} ({study.patient.patient_id})")
    doc.add_paragraph(f"Accession: {study.accession_number}    Modality: {study.modality.code}    Date: {study.study_date}")
    doc.add_paragraph(f"Priority: {study.priority.upper()}")
    doc.add_paragraph('')
    sections = [
        ('Clinical History', (report.clinical_history if report else (study.clinical_info or '')) or '-'),
        ('Technique', (report.technique if report else '') or '-'),
        ('Comparison', (report.comparison if report else '') or '-'),
        ('Findings', (report.findings if report else '') or '-'),
        ('Impression', (report.impression if report else '') or '-'),
        ('Recommendations', (report.recommendations if report else '') or '-'),
    ]
    for title, content in sections:
        doc.add_heading(title, level=2)
        doc.add_paragraph(content)
    # Signature
    author_name = ''
    author_license = ''
    if report and getattr(report, 'radiologist', None):
        try:
            author_name = report.radiologist.get_full_name() or report.radiologist.username
            author_license = getattr(report.radiologist, 'license_number', '') or ''
        except Exception:
            pass
    doc.add_paragraph(f"Signed by: {author_name}{(' - ' + author_license) if author_license else ''}")
    if report and report.signed_date:
        doc.add_paragraph(f"Signed on: {report.signed_date.strftime('%Y-%m-%d %H:%M')}")
    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    resp = HttpResponse(buf.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    resp['Content-Disposition'] = f'attachment; filename="report_{slugify(study.accession_number)}.docx"'
    return resp
